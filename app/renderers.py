from __future__ import annotations

from dataclasses import dataclass
import json
from pathlib import Path
from typing import Iterable


@dataclass(frozen=True)
class TranscriptSegment:
    start: float
    end: float
    text: str


@dataclass(frozen=True)
class TranscriptParagraph:
    start: float
    end: float
    text: str


def _clean_text(text: str) -> str:
    return " ".join(str(text or "").split()).strip()


def format_timestamp(seconds: float) -> str:
    total = max(0, int(seconds))
    hours = total // 3600
    minutes = (total % 3600) // 60
    secs = total % 60
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def merge_segments(segments: list[TranscriptSegment]) -> list[TranscriptParagraph]:
    paragraphs: list[TranscriptParagraph] = []
    current_texts: list[str] = []
    current_start = 0.0
    current_end = 0.0

    for segment in segments:
        text = _clean_text(segment.text)
        if not text:
            continue

        if not current_texts:
            current_start = segment.start
            current_end = segment.end
            current_texts = [text]
            continue

        previous_text = current_texts[-1]
        gap = max(0.0, segment.start - current_end)
        should_split = (
            gap >= 1.6
            or len(" ".join(current_texts)) >= 120
            or previous_text.endswith(("。", "！", "？", ".", "!", "?"))
        )
        if should_split:
            paragraphs.append(
                TranscriptParagraph(
                    start=current_start,
                    end=current_end,
                    text=" ".join(current_texts).strip(),
                )
            )
            current_start = segment.start
            current_texts = [text]
        else:
            current_texts.append(text)

        current_end = segment.end

    if current_texts:
        paragraphs.append(
            TranscriptParagraph(
                start=current_start,
                end=current_end,
                text=" ".join(current_texts).strip(),
            )
        )

    return paragraphs


def _meta_lines(title: str, source_url: str, language: str, duration_seconds: float) -> list[str]:
    return [
        f"标题：{title}",
        f"来源：{source_url}" if source_url else "",
        f"语言：{language}" if language else "",
        f"时长：{format_timestamp(duration_seconds)}" if duration_seconds else "",
    ]


def _render_txt(title: str, source_url: str, language: str, duration_seconds: float, paragraphs: list[TranscriptParagraph]) -> str:
    lines = [line for line in _meta_lines(title, source_url, language, duration_seconds) if line]
    lines.append("")
    for paragraph in paragraphs:
        lines.append(paragraph.text)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _render_markdown(title: str, source_url: str, language: str, duration_seconds: float, paragraphs: list[TranscriptParagraph]) -> str:
    lines = [f"# {title}", ""]
    for line in _meta_lines(title, source_url, language, duration_seconds):
        if line:
            lines.append(f"- {line}")
    lines.append("")
    lines.append("## 正文")
    lines.append("")
    for paragraph in paragraphs:
        lines.append(paragraph.text)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _render_script_markdown(title: str, segments: list[TranscriptSegment]) -> str:
    lines = [f"# {title} 台本", ""]
    for segment in segments:
        text = _clean_text(segment.text)
        if not text:
            continue
        lines.append(f"- `{format_timestamp(segment.start)} - {format_timestamp(segment.end)}` {text}")
    return "\n".join(lines).strip() + "\n"


def _render_json(title: str, source_url: str, language: str, duration_seconds: float, segments: list[TranscriptSegment], paragraphs: list[TranscriptParagraph]) -> str:
    payload = {
        "title": title,
        "source_url": source_url,
        "language": language,
        "duration_seconds": duration_seconds,
        "segments": [
            {
                "start": segment.start,
                "end": segment.end,
                "text": _clean_text(segment.text),
            }
            for segment in segments
        ],
        "paragraphs": [
            {
                "start": paragraph.start,
                "end": paragraph.end,
                "text": paragraph.text,
            }
            for paragraph in paragraphs
        ],
    }
    return json.dumps(payload, ensure_ascii=False, indent=2) + "\n"


def _write_docx(path: Path, title: str, source_url: str, language: str, duration_seconds: float, paragraphs: list[TranscriptParagraph]) -> None:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt

    document = Document()
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)

    for line in _meta_lines(title, source_url, language, duration_seconds):
        if not line:
            continue
        meta_paragraph = document.add_paragraph()
        meta_paragraph.paragraph_format.space_after = Pt(2)
        meta_run = meta_paragraph.add_run(line)
        meta_run.font.size = Pt(9.5)

    for paragraph in paragraphs:
        body = document.add_paragraph()
        body.paragraph_format.first_line_indent = Pt(24)
        body.paragraph_format.space_before = Pt(6)
        body.paragraph_format.space_after = Pt(6)
        body.paragraph_format.line_spacing = 1.45
        run = body.add_run(paragraph.text)
        run.font.size = Pt(11.5)

    document.save(path)


def normalize_output_formats(formats: list[str] | tuple[str, ...] | set[str] | None) -> list[str]:
    ordered: list[str] = []
    for item in list(formats or ["txt"]):
        value = str(item).strip().lower()
        if not value:
            continue
        if value == "word":
            value = "docx"
        if value not in {"txt", "md", "docx"}:
            continue
        if value not in ordered:
            ordered.append(value)
    if "txt" not in ordered:
        ordered.insert(0, "txt")
    return ordered


def write_transcript_bundle(
    output_dir: Path,
    title: str,
    source_url: str,
    language: str,
    duration_seconds: float,
    segments: list[TranscriptSegment],
    *,
    paragraphs: list[TranscriptParagraph] | None = None,
    output_formats: list[str] | tuple[str, ...] | set[str] | None = None,
) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    normalized_paragraphs = paragraphs or merge_segments(segments)
    selected_formats = normalize_output_formats(output_formats)
    paths: list[Path] = []

    if "txt" in selected_formats:
        txt_path = output_dir / "transcript.txt"
        txt_path.write_text(
            _render_txt(title, source_url, language, duration_seconds, normalized_paragraphs),
            encoding="utf-8",
        )
        paths.append(txt_path)

    if "md" in selected_formats:
        md_path = output_dir / "transcript.md"
        md_path.write_text(
            _render_markdown(title, source_url, language, duration_seconds, normalized_paragraphs),
            encoding="utf-8",
        )
        paths.append(md_path)

    if "docx" in selected_formats:
        docx_path = output_dir / "transcript.docx"
        _write_docx(docx_path, title, source_url, language, duration_seconds, normalized_paragraphs)
        paths.append(docx_path)

    return {
        "paragraph_count": len(normalized_paragraphs),
        "segment_count": len(segments),
        "paths": paths,
    }
